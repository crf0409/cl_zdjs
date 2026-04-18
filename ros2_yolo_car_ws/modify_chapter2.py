#!/usr/bin/env python3
"""修改第二章：去ROS → 3D语义占用预测技术"""
import copy
import os
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

DOCX_PATH = '/home/siton02/md0/crf/ros2_yolo_car_ws/第二章_相关技术与理论基础.docx'
OUTPUT_PATH = '/home/siton02/md0/crf/ros2_yolo_car_ws/第二章_相关技术与理论基础_modified.docx'
FIGURES_DIR = '/home/siton02/md0/crf/ros2_yolo_car_ws/figures'


def delete_paragraph(para):
    """Delete a paragraph from the document."""
    p = para._element
    p.getparent().remove(p)


def insert_paragraph_after(doc, ref_para, text, style='Normal'):
    """Insert a new paragraph after ref_para with given text and style."""
    new_p = copy.deepcopy(ref_para._element)
    # Clear all runs
    for r in new_p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
        new_p.remove(r)
    # Clear drawings/images
    for drawing in new_p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
        new_p.remove(drawing)

    # Create new run with text
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    new_r = etree.SubElement(new_p, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
    new_t = etree.SubElement(new_r, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
    new_t.text = text
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    ref_para._element.addnext(new_p)
    # Return the new paragraph object
    # Find it in doc.paragraphs
    for p in doc.paragraphs:
        if p._element is new_p:
            return p
    return None


def add_paragraph_with_style(doc, ref_para, text, heading_level=None):
    """Add paragraph after ref_para, copying format from a same-level heading or normal text."""
    # Find a reference paragraph with the desired style
    if heading_level:
        style_name = f'Heading {heading_level}'
    else:
        style_name = 'Normal'

    # Find a paragraph with this style to copy from
    template = None
    for p in doc.paragraphs:
        if p.style and p.style.name == style_name and p.text.strip():
            template = p
            break

    if template is None:
        template = ref_para

    new_p = copy.deepcopy(template._element)
    # Clear runs
    for r in new_p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
        new_p.remove(r)
    # Clear drawings
    for d in new_p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
        new_p.remove(d)
    for d in new_p.findall('.//{http://schemas.microsoft.com/office/word/2010/wordml}drawing'):
        new_p.remove(d)

    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    new_r = etree.SubElement(new_p, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')

    # Copy run properties from template if available
    if template.runs:
        rpr = template.runs[0]._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
        if rpr is not None:
            new_r.insert(0, copy.deepcopy(rpr))

    new_t = etree.SubElement(new_r, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
    new_t.text = text
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    ref_para._element.addnext(new_p)

    for p in doc.paragraphs:
        if p._element is new_p:
            return p
    return None


def insert_image_paragraph(doc, ref_para, image_path, caption_text, width_inches=5.5):
    """Insert an image and caption after ref_para."""
    # Insert caption first, then image before it
    cap = add_paragraph_with_style(doc, ref_para, caption_text)
    if os.path.exists(image_path):
        # Add image to the paragraph before caption
        img_p = add_paragraph_with_style(doc, ref_para, '')
        # We need to add the image directly
        run = img_p.runs[0] if img_p and img_p.runs else None
        if run:
            run.clear()
        # Alternative: use doc level to add picture (append only)
        # For now, use placeholder text
        if img_p:
            for r in img_p._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
                for t in r.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                    t.text = f'[此处插入图片：{image_path}]'
        return cap
    else:
        # Image doesn't exist, use placeholder
        placeholder = add_paragraph_with_style(doc, ref_para, f'[此处插入图片：{os.path.basename(image_path)}]')
        return cap


def replace_text_in_paragraph(para, old_text, new_text):
    """Replace text in a paragraph while preserving formatting."""
    for run in para.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)


def set_paragraph_text(para, new_text):
    """Replace all text in a paragraph."""
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ''
    else:
        # No runs, create one
        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        new_r = etree.SubElement(para._element, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
        new_t = etree.SubElement(new_r, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
        new_t.text = new_text


def main():
    doc = Document(DOCX_PATH)
    paras = doc.paragraphs

    print(f"Original paragraphs: {len(paras)}")

    # =========================================================
    # STEP 1: Modify chapter intro (para 1)
    # =========================================================
    set_paragraph_text(paras[1],
        '本章系统介绍本课题涉及的关键技术与理论基础，包括YOLO目标检测算法原理、'
        '3D语义占用预测技术、常用数据集，以及模型轻量化与嵌入式部署技术。'
        '这些技术共同构成了基于视觉的智能障碍物感知系统的理论支撑。')

    # =========================================================
    # STEP 2: Delete section 2.2 ROS (paras 20-28)
    # Delete from end to start to preserve indices
    # =========================================================
    print("Deleting section 2.2 (ROS)...")
    ros_paras = list(paras[20:29])  # paras 20-28
    for p in reversed(ros_paras):
        delete_paragraph(p)

    # Refresh paragraph list
    paras = doc.paragraphs
    print(f"After deleting ROS section: {len(paras)} paragraphs")

    # =========================================================
    # STEP 3: Insert new section 2.2 (3D occupancy prediction)
    # Insert after para 19 (last para of section 2.1)
    # =========================================================
    print("Inserting new section 2.2 (3D occupancy prediction)...")

    # Current para 19 is the last paragraph of section 2.1
    anchor = paras[19]

    new_section_content = [
        (2, '2.2 3D语义占用预测技术'),
        (3, '2.2.1 BEV感知与Lift-Splat-Shoot视角变换'),
        (None, 'BEV（Bird\'s Eye View，鸟瞰视角）感知是自动驾驶领域近年来最重要的技术范式之一。'
         '与传统的2D图像检测不同，BEV感知将多个摄像头的2D图像特征统一变换到3D空间的鸟瞰视角表示中，'
         '从而获得对周围环境的全局化、结构化理解。这种表示方式天然适合自动驾驶中的障碍物检测、'
         '路径规划和运动预测等下游任务。'),
        (None, 'Lift-Splat-Shoot（LSS）是由Philion和Fidler于2020年提出的经典BEV视角变换框架，'
         '其核心思想是通过显式的深度估计将2D图像特征"提升"到3D空间。具体而言，LSS包含三个阶段：'
         '（1）Lift（提升）——对每个图像像素预测离散的深度概率分布，将2D特征沿深度方向展开形成3D视锥体特征；'
         '（2）Splat（投射）——将3D视锥体特征投影到统一的BEV网格中，通过柱体求和完成多视角特征的融合；'
         '（3）Shoot（编码）——对BEV特征进行进一步的卷积编码，用于下游的语义分割或目标检测任务。'
         'LSS框架的核心优势在于仅依赖摄像头输入即可构建3D场景表示，无需昂贵的LiDAR传感器。'),
        (None, '[此处插入图2-3：Lift-Splat-Shoot视角变换流程示意图]'),
        (None, '图2-3 Lift-Splat-Shoot视角变换流程示意图'),
        (3, '2.2.2 3D语义占用预测任务定义'),
        (None, '3D语义占用预测（3D Semantic Occupancy Prediction）是BEV感知的进一步延伸。'
         '该任务的目标是：给定多视角摄像头图像，预测一个密集的3D体素网格，其中每个体素被标注为'
         '特定的语义类别（如汽车、行人、可行驶路面、植被等）或"空闲空间"。'
         '相比传统的2D边界框检测，3D占用预测具有三大优势：'
         '（1）提供对场景的体积化理解，能够感知任意形状的障碍物（包括不规则形状的护栏、路缘石等）；'
         '（2）显式地建模空闲空间，为路径规划提供安全通行区域信息；'
         '（3）能够处理遮挡区域，通过学习到的3D先验推断被前方物体遮挡的场景结构。'),
        (None, '在Occ3D-nuScenes基准中，占用网格的规格为200×200×16体素，'
         '覆盖自车周围80m×80m×6.4m的空间范围（体素分辨率0.4m）。'
         '语义类别共19类，包括10类前景目标（汽车、卡车、工程车、公交车、拖车、护栏、'
         '摩托车、自行车、行人、锥桶）和6类背景（可行驶路面、其他地面、人行道、地形、'
         '人造物、植被），以及空闲空间和噪声类。评估指标采用平均交并比mIoU（mean Intersection over Union），'
         '即各类别IoU的算术平均值。'),
        (3, '2.2.3 ALOcc框架原理'),
        (None, 'ALOcc（Adaptive Lifting-based 3D Semantic Occupancy）是2025年发表于ICCV的'
         '最新3D占用预测框架。ALOcc采用纯卷积设计，通过自适应提升机制将多视角2D图像特征高效地'
         '变换为3D体素表示，在精度和速度之间取得了出色的平衡。'),
        (None, 'ALOcc的完整推理流程如下：'
         '（1）图像骨干网络（Image Backbone）——本课题采用YOLOv8s的CSPDarknet作为特征提取器，'
         '替换原始的ResNet50，输出P3（步长8）、P4（步长16）、P5（步长32）三个尺度的特征图；'
         '（2）特征金字塔网络（FPN Neck）——将P4和P5特征进行双向融合，输出256通道的统一特征；'
         '（3）深度估计网络（CM_DepthNet）——预测每个像素的离散深度概率分布（88个深度bin，覆盖1.0~45.0m）'
         '及80通道的上下文特征，采用ASPP（空洞空间金字塔池化）模块增强多尺度感受野；'
         '（4）LSS视角变换器——基于预测的深度分布将2D特征提升为3D视锥体，投射到200×200×16的BEV/体素网格中；'
         '（5）BEV编码器——由4阶段CustomResNet（通道数[64,128,256,512]）和FPN_LSS2构成，'
         '对BEV特征进行多尺度编码和融合；'
         '（6）ALOcc Head——基于Mask2Former架构，使用19个可学习的类别查询（class queries），'
         '通过交叉注意力机制与BEV特征交互，最终预测每个体素的语义类别。'),
        (None, '[此处插入图2-4：ALOcc+YOLOv8s系统架构示意图]'),
        (None, '图2-4 ALOcc+YOLOv8s系统架构示意图'),
        (None, '在本课题中，将YOLOv8s的CSPDarknet骨干网络替换ALOcc原始的ResNet50，'
         '通过1×1卷积+BN+ReLU的通道适配层将YOLOv8s的输出通道（P3:128→256, P4:256→1024, P5:512→2048）'
         '映射到ALOcc FPN所期望的通道维度。这种设计使得同一个YOLO骨干网络既可用于独立的2D目标检测，'
         '也可作为ALOcc 3D占用预测的特征提取前端，实现"一个骨干、双重任务"的高效架构。'),
    ]

    last_inserted = anchor
    for heading_level, text in new_section_content:
        last_inserted = add_paragraph_with_style(doc, last_inserted, text, heading_level=heading_level)

    paras = doc.paragraphs
    print(f"After inserting 2.2: {len(paras)} paragraphs")

    # =========================================================
    # STEP 4: Add Occ3D content to nuScenes section (2.3.1)
    # Find "图2-5展示了nuScenes" paragraph and add after it
    # =========================================================
    print("Adding Occ3D content to nuScenes section...")
    for i, p in enumerate(paras):
        if '本课题从nuScenes mini-val子集中提取81帧' in p.text or '图2-5展示了nuScenes数据集中的四种典型驾驶场景' in p.text:
            occ3d_text = (
                '除了2D目标检测标注外，nuScenes数据集还配备了Occ3D-nuScenes 3D占用标注。'
                'Occ3D标注通过累积多帧LiDAR点云并进行体素化处理生成，为每个体素分配19类语义标签。'
                '占用网格覆盖[-40m, 40m]×[-40m, 40m]×[-1m, 5.4m]的3D空间，分辨率为0.4m，'
                '网格尺寸为200×200×16。这些密集的3D标注为训练和评估ALOcc等3D占用预测模型提供了'
                '高质量的监督信号，是本课题进行3D场景理解实验的核心数据基础。'
            )
            add_paragraph_with_style(doc, p, occ3d_text)
            break

    paras = doc.paragraphs

    # =========================================================
    # STEP 5: Simplify Gazebo section (2.3.2)
    # Find the Gazebo heading and replace content
    # =========================================================
    print("Simplifying Gazebo section...")
    gazebo_start = None
    gazebo_end = None
    for i, p in enumerate(paras):
        if p.text.strip() == '2.3.2 Gazebo仿真平台介绍':
            gazebo_start = i
        elif gazebo_start and p.style and 'Heading' in p.style.name and i > gazebo_start:
            gazebo_end = i
            break

    if gazebo_start and gazebo_end:
        # Delete paragraphs between heading and next heading (paras gazebo_start+1 to gazebo_end-1)
        to_delete = list(paras[gazebo_start + 1:gazebo_end])
        for p in reversed(to_delete):
            delete_paragraph(p)

        paras = doc.paragraphs
        # Insert simplified Gazebo text after the heading
        for i, p in enumerate(paras):
            if p.text.strip() == '2.3.2 Gazebo仿真平台介绍':
                # Rename to simpler title
                set_paragraph_text(p, '2.3.2 Gazebo仿真平台简介')
                simplified = (
                    'Gazebo是一个高保真的机器人仿真平台，能够模拟复杂的三维物理环境和多种传感器模型。'
                    '本课题使用Gazebo Harmonic作为辅助验证工具，构建TurtleBot3 World虚拟场景，'
                    '配备RGB摄像头（640×480, 30Hz）和360度激光雷达，用于对避障决策逻辑进行可视化验证。'
                    'Gazebo在本课题中起补充作用，主要实验基于nuScenes真实驾驶数据集进行。'
                )
                add_paragraph_with_style(doc, p, simplified)
                break

    paras = doc.paragraphs

    # =========================================================
    # STEP 6: Rewrite 2.3.3 application value
    # =========================================================
    print("Rewriting 2.3.3 application value...")
    for i, p in enumerate(paras):
        if '2.3.3' in p.text and '应用价值' in p.text:
            set_paragraph_text(p, '2.3.3 nuScenes数据集在本课题中的应用价值')
            # Find and replace the content paragraphs
            j = i + 1
            while j < len(paras) and (not paras[j].style or 'Heading' not in paras[j].style.name):
                if paras[j].text.strip():
                    set_paragraph_text(paras[j],
                        'nuScenes数据集在本课题中承担双重角色：'
                        '（1）其2D标注图像（1600×900前视摄像头）为YOLO模型的2D障碍物检测评估提供测试数据——'
                        '本文从mini-val子集提取81帧图像，使用COCO预训练的YOLO模型进行推理评估；'
                        '（2）其Occ3D 3D占用标注为ALOcc框架的3D语义占用预测训练与评估提供监督信号——'
                        '本文使用mini训练集（323样本）进行快速验证实验，完整训练集（28130样本）可用于大规模训练。'
                        '这种"2D检测+3D占用"的双任务评估策略，使得同一个YOLO骨干网络的能力可以从'
                        '不同维度得到全面验证。'
                    )
                    # Delete the second paragraph if exists
                    if j + 1 < len(paras) and paras[j + 1].text.strip() and 'Heading' not in (paras[j + 1].style.name if paras[j + 1].style else ''):
                        if '双数据源策略' in paras[j + 1].text or '仿真数据' in paras[j + 1].text:
                            delete_paragraph(paras[j + 1])
                    break
                j += 1
            break

    paras = doc.paragraphs

    # =========================================================
    # STEP 7: Simplify Jetson section (2.4.2)
    # =========================================================
    print("Simplifying Jetson section...")
    for i, p in enumerate(paras):
        if '2.4.2' in p.text and 'Jetson' in p.text:
            # Find content paragraphs and simplify
            j = i + 1
            first_content = True
            while j < len(paras) and (not paras[j].style or 'Heading' not in paras[j].style.name):
                if paras[j].text.strip() and first_content:
                    # Keep first content paragraph but simplify
                    if '图2-9' in paras[j].text:
                        # This is a figure caption, skip
                        j += 1
                        continue
                    set_paragraph_text(paras[j],
                        'Jetson Nano配备128核Maxwell GPU和4GB共享内存，AI算力为472 GFLOPS，'
                        '是移动机器人部署的高性价比选择。本课题的开发与验证阶段使用配备RTX 3090的工作站'
                        '（Ubuntu 24.04, CUDA 12.4）运行mmdet3d框架进行ALOcc模型训练与评估。'
                        '对于独立的YOLO 2D检测任务，经TensorRT INT8量化后可在Jetson Nano上以15~18 FPS运行；'
                        'ALOcc完整推理管线因参数量较大（48.4M），目前以工作站GPU为目标平台。'
                    )
                    first_content = False
                elif not first_content and paras[j].text.strip():
                    # Delete extra paragraphs (Xavier NX description etc.)
                    if 'Xavier NX' in paras[j].text:
                        delete_paragraph(paras[j])
                        continue
                j += 1
            break

    paras = doc.paragraphs

    # =========================================================
    # STEP 8: Rewrite chapter summary (2.5)
    # =========================================================
    print("Rewriting chapter summary...")
    for i, p in enumerate(paras):
        if p.text.strip() == '2.5 本章小结':
            # Replace all following paragraphs until end
            j = i + 1
            summary_written = False
            while j < len(paras):
                if not summary_written:
                    set_paragraph_text(paras[j],
                        '本章系统介绍了本课题涉及的四项关键技术：'
                        '（1）YOLO目标检测算法从YOLOv1到YOLO11的演进历程和核心原理，'
                        '以及YOLOv5与YOLO11在网络结构、检测头设计和训练策略方面的主要差异。'
                        'YOLO的高速度、高精度和多规模档位设计使其既适合作为独立的2D检测器，'
                        '也适合作为3D感知框架的骨干网络。'
                    )
                    summary_written = True
                    j += 1
                    if j < len(paras):
                        set_paragraph_text(paras[j],
                            '（2）3D语义占用预测技术，包括BEV感知范式、'
                            'Lift-Splat-Shoot视角变换框架和ALOcc占用预测框架。'
                            'ALOcc通过自适应提升机制将多视角2D图像特征变换为200×200×16的3D语义体素网格，'
                            '基于Mask2Former架构的ALOcc Head实现19类语义的精细预测。'
                            '本课题创新性地采用YOLOv8s替换原始ResNet50作为ALOcc的图像骨干网络，'
                            '实现了"一个骨干、双重任务"的高效架构设计。'
                        )
                    j += 1
                    if j < len(paras):
                        set_paragraph_text(paras[j],
                            '（3）nuScenes自动驾驶数据集及其Occ3D 3D占用标注，'
                            '前者提供多样化的真实驾驶场景用于2D检测评估，'
                            '后者提供密集的3D语义体素标注用于占用预测训练与评估。'
                            'Gazebo仿真平台作为辅助工具用于避障决策的可视化验证。'
                        )
                    j += 1
                    if j < len(paras):
                        set_paragraph_text(paras[j],
                            '（4）TensorRT推理优化通过层融合、精度量化和内核调优等技术，'
                            '可将YOLO模型的推理速度提升2.5~3.5倍。'
                            'Jetson Nano等嵌入式平台适用于YOLO独立2D检测的实时部署，'
                            'ALOcc完整管线以工作站GPU为目标平台。'
                            '这些技术基础为后续章节的模型设计、系统部署和实验验证奠定了理论支撑。'
                        )
                    break
                j += 1
            break

    paras = doc.paragraphs

    # =========================================================
    # STEP 9: Update figure references throughout the document
    # Old 图2-4 → 图2-5, old 图2-5 → 图2-6
    # (图2-3 ROS was deleted, new 图2-3 and 图2-4 were inserted)
    # =========================================================
    print("Updating figure references...")
    # First rename old references to avoid conflicts (use temporary names)
    renames = [
        ('图2-5展示了nuScenes', '图2-6展示了nuScenes'),
        ('图2-5 nuScenes典型场景', '图2-6 nuScenes典型场景'),
        ('如图2-4所示', '如图2-5所示'),
        ('图2-4 nuScenes自动驾驶', '图2-5 nuScenes自动驾驶'),
        ('图2-7对比了', '图2-7对比了'),  # stays the same
        ('图2-7 常用目标检测', '图2-7 常用目标检测'),  # stays
    ]

    for p in paras:
        for old, new in renames:
            if old in p.text and old != new:
                for run in p.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)

    # =========================================================
    # STEP 10: Save
    # =========================================================
    doc.save(OUTPUT_PATH)
    print(f"\nSaved to: {OUTPUT_PATH}")

    # Print final structure
    print("\n=== Final Structure ===")
    for i, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style else ''
        text = p.text.strip()[:80] if p.text.strip() else ''
        if 'Heading' in style or text:
            prefix = f'[{style}] ' if 'Heading' in style else '  '
            print(f'{i:3d} {prefix}{text}')


if __name__ == '__main__':
    main()
