#!/usr/bin/env python3
"""生成第二章 相关技术与理论基础 Word文档"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT_DOCX = '/home/siton02/md0/ros2_yolo_car_ws/第二章_相关技术与理论基础.docx'
FIG = '/home/siton02/md0/ros2_yolo_car_ws/chapter2_figures'
YOLO_RES = '/home/siton02/md0/crf/cl_zdjs/project1_yolo_obstacle/results'

doc = Document()
style = doc.styles['Normal']
style.font.name = '宋体'; style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.first_line_indent = Cm(0.74)

def H(text, level):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0,0,0); r.font.name='黑体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
        r.font.size = Pt([0,16,14,12][level])
    h.paragraph_format.first_line_indent = Cm(0)

def B(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    for r in p.runs:
        r.font.name='宋体'; r.font.size=Pt(12)
        r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

def F(path, cap, w=5.8):
    if not os.path.exists(path):
        B(f'[图片缺失: {os.path.basename(path)}]'); return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run().add_picture(path, width=Inches(w))
    c = doc.add_paragraph(cap)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.first_line_indent = Cm(0)
    c.paragraph_format.space_after = Pt(6)
    for r in c.runs: r.font.size=Pt(10); r.font.name='宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

# ════════════════════════════════════════════════════
H('第二章 相关技术与理论基础', 1)

B('本章系统介绍本课题涉及的关键技术与理论基础，包括YOLO目标检测算法原理、'
  'ROS机器人操作系统架构、常用数据集与仿真平台，以及模型轻量化与嵌入式部署技术。'
  '这些技术共同构成了基于视觉的移动机器人智能避障系统的理论支撑。')

# ── 2.1 ──
H('2.1 YOLO目标检测算法原理', 2)

H('2.1.1 YOLO系列发展历程与核心思想', 3)

B('YOLO（You Only Look Once）算法是目前最具影响力的单阶段目标检测框架之一。'
  '与传统的两阶段检测算法（如Faster R-CNN先生成候选区域再分类）不同，'
  'YOLO将目标检测建模为一个端到端的回归问题——'
  '仅通过一次网络前向传播即可同时输出所有目标的位置、类别和置信度。'
  '这种设计从根本上消除了候选区域生成的计算开销，'
  '使得YOLO系列算法在推理速度上具有显著优势。')

B('YOLO系列经历了多代演进。YOLOv1（2016年）首次提出将检测问题转化为回归问题，'
  '在速度上实现了45 FPS的突破但精度有限。YOLOv2/v3分别引入了BN层、多尺度检测和'
  'Darknet-53骨干网络，精度大幅提升。YOLOv4（2020年）集成了CSPDarknet、'
  'Mish激活函数和Mosaic数据增强等技术，被认为是该系列的里程碑版本。'
  'YOLOv5由Ultralytics公司基于PyTorch重新实现，提供了n/s/m/l/x五个规模档位，'
  '极大降低了使用和部署门槛。YOLOv8/YOLO11延续Ultralytics框架，'
  '引入解耦头、C2f模块和无锚框设计，在COCO数据集上的mAP达到44.9（s规模），'
  '相比YOLOv5s的37.4提升了20%。')

B('YOLO的核心网络架构由三部分组成：Backbone（骨干网络）负责从输入图像中提取'
  '多尺度特征；Neck（特征融合模块）通过FPN+PAN结构将不同尺度的特征进行双向融合；'
  'Head（检测头）输出最终的边界框坐标、目标类别概率和置信度得分。')

F(os.path.join(FIG, 'fig_2_1_yolo_architecture.png'),
  '图2-1 YOLO目标检测网络架构示意图')

B('如图2-1所示，输入图像经过Backbone的逐层卷积提取出三个尺度的特征图'
  '（P3/P4/P5，分辨率分别为80×80、40×40、20×20），'
  '分别对应小、中、大目标的检测。Neck模块将高层语义特征和低层纹理特征'
  '进行双向融合——自顶向下的FPN路径将深层语义信息传递给浅层，'
  '自底向上的PAN路径将浅层的精确定位信息传递给深层。'
  '融合后的多尺度特征送入检测头进行最终的分类与回归预测。')

H('2.1.2 YOLOv5与YOLO11网络结构对比', 3)

B('YOLOv5和YOLO11（YOLOv8的后续演进版本）是目前工程部署中最常用的两个YOLO版本。'
  '二者均采用CSPDarknet骨干网络，但在模块设计、检测头结构和训练策略等方面存在显著差异。')

F(os.path.join(FIG, 'fig_2_2_v5_vs_v11_comparison.png'),
  '图2-2 YOLOv5与YOLO11关键技术差异对比')

B('如图2-2所示，两个版本的核心差异体现在四个方面：'
  '（1）骨干网络中，YOLO11用C2f模块替代了YOLOv5的C3模块。'
  'C2f通过分割-拼接的方式增加了梯度流的多样性，在不增加计算量的前提下提升了特征提取效率。'
  '（2）检测头方面，YOLO11采用解耦头（Decoupled Head）将分类和回归分支完全分离，'
  '避免了耦合头中两个任务相互干扰的问题，这对小目标检测尤为重要。'
  '（3）锚框策略方面，YOLO11彻底抛弃了预定义锚框，采用Anchor-free设计，'
  '直接预测目标中心点到边界框四边的距离，简化了超参数设计。'
  '（4）标签分配方面，YOLO11使用TaskAligned Assigner替代了YOLOv5的IoU匹配策略，'
  '根据分类和回归的联合得分动态分配正样本，训练更加稳定。')

H('2.1.3 YOLO在移动机器人障碍物检测中的适用性分析', 3)

B('移动机器人避障场景对目标检测算法提出了三项核心需求：'
  '实时性（检测频率≥10Hz）、准确性（不漏检近距离障碍物）'
  '和轻量化（适应嵌入式算力约束）。YOLO系列算法在这三个维度上均展现了良好的适配性。')

B('在实时性方面，YOLO11s在RTX 3090上的推理帧率达到43.3 FPS，'
  '远超10 Hz的最低要求。即便在Jetson Nano等边缘设备上，'
  '经TensorRT INT8量化后仍可达到15～18 FPS，满足低速移动机器人的实时性需求。')

B('在准确性方面，YOLO的多尺度检测机制使其能同时感知远处的小目标'
  '（提供早期预警）和近处的大目标（触发紧急避障）。'
  'COCO预训练权重已覆盖80类常见目标，无需额外训练即可检测大部分障碍物。')

B('在轻量化方面，YOLO提供n/s/m/l多个规模档位，'
  '用户可根据算力预算灵活选择。YOLO11s仅有9.4M参数（模型文件约19MB），'
  '可轻松部署在4GB内存的Jetson Nano上。')

# ── 2.2 ──
H('2.2 机器人操作系统（ROS）基础', 2)

H('2.2.1 ROS节点、话题、服务的架构', 3)

B('ROS（Robot Operating System）是一个面向机器人应用的开源软件框架，'
  '提供了节点管理、消息通信、包管理和工具链等基础设施。'
  'ROS 2作为新一代版本，采用DDS（Data Distribution Service）中间件替代了ROS 1的'
  '自定义TCP通信协议，在实时性、可靠性和安全性方面有显著提升。')

F(os.path.join(FIG, 'fig_2_3_ros2_architecture.png'),
  '图2-3 ROS 2节点-话题-服务通信架构')

B('如图2-3所示，ROS 2的通信模型包含三种核心机制：'
  '（1）话题（Topic）：基于发布/订阅模式的异步通信，适用于连续数据流（如相机图像、检测结果）。'
  '发布者和订阅者通过话题名称解耦，支持一对多和多对一通信。'
  '（2）服务（Service）：基于请求/响应模式的同步通信，适用于一次性查询操作。'
  '（3）动作（Action）：带有反馈的长时间异步任务，适用于导航等耗时操作。'
  '所有通信均建立在DDS中间件之上，支持QoS（服务质量）策略配置、'
  '自动节点发现和跨进程/跨机器通信。')

H('2.2.2 ROS与视觉感知、运动控制的集成方式', 3)

B('在本课题的避障系统中，ROS 2承担了视觉感知与运动控制的集成枢纽角色。'
  '相机驱动节点（或Gazebo仿真节点）发布sensor_msgs/Image类型的图像消息至'
  '/camera/image_raw话题；YOLO检测节点订阅该话题，通过cv_bridge库将ROS图像消息'
  '转换为OpenCV格式后进行GPU推理，检测结果以vision_msgs/Detection2DArray格式发布；'
  '避障决策节点订阅检测结果，应用三区域反应式算法生成geometry_msgs/TwistStamped'
  '速度指令发布至/cmd_vel话题，由底层驱动转化为电机控制信号。')

B('ROS 2的话题通信机制天然适合这种流水线式的数据处理架构——'
  '各节点独立运行、松耦合连接，任一节点的更新或替换不影响其他节点。'
  '例如，可以在不修改避障节点的前提下将YOLO11s替换为YOLO11m，'
  '或将仿真相机替换为实际USB摄像头。')

# ── 2.3 ──
H('2.3 障碍物检测常用数据集与仿真平台', 2)

H('2.3.1 nuScenes数据集介绍', 3)

B('nuScenes是由Motional公司发布的大规模自动驾驶数据集，'
  '于2019年首次公开。该数据集包含1000个驾驶场景（每个场景约20秒），'
  '覆盖波士顿和新加坡两座城市的多种道路环境。'
  '数据采集车辆配备了全套多模态传感器：6个环视摄像头（1600×900分辨率）、'
  '1个32线激光雷达和5个毫米波雷达，实现了360度全方位感知覆盖。')

F(os.path.join(FIG, 'fig_2_5_nuscenes_sensors.png'),
  '图2-4 nuScenes自动驾驶平台传感器配置示意图', 4.5)

B('如图2-4所示，6个摄像头分别朝向前方、前左、前右、后方、后左和后右六个方向，'
  '无盲区地覆盖了车辆周围的完整视野。32线激光雷达安装于车顶，'
  '提供高精度的3D点云数据。5个毫米波雷达分布在车身四周，提供远距离测速信息。'
  '这种多传感器配置使nuScenes成为研究多模态融合感知的理想平台。')

F(os.path.join(FIG, 'fig_2_4_nuscenes_overview.png'),
  '图2-5 nuScenes典型场景示例（前视摄像头）')

B('图2-5展示了nuScenes数据集中的四种典型驾驶场景：城市桥梁、城市街道、'
  '十字路口和居民区。场景多样性涵盖了不同的道路结构、交通密度和光照条件，'
  '为目标检测算法的泛化能力评估提供了丰富的测试样本。'
  '本课题从nuScenes mini-val子集中提取81帧前视图像作为离线评估数据集，'
  '这些图像包含车辆、行人、自行车、交通灯等多种避障相关目标。')

H('2.3.2 Gazebo仿真平台介绍', 3)

B('Gazebo是一个高保真的机器人仿真平台，能够模拟复杂的三维物理环境、'
  '多种传感器模型以及机器人的运动学和动力学行为。'
  'Gazebo Harmonic（8.x）是其最新长期支持版本，'
  '原生支持ROS 2 Jazzy，通过ros_gz_bridge实现Gazebo话题与ROS 2话题的双向桥接。')

F(os.path.join(FIG, 'fig_2_6_gazebo_ros2_arch.png'),
  '图2-6 Gazebo Harmonic与ROS 2集成架构')

B('如图2-6所示，Gazebo内部由物理引擎（DART/ODE）、渲染引擎（Ogre2）、'
  '传感器模型和世界模型四大模块构成。物理引擎负责刚体碰撞和摩擦力计算，'
  '确保机器人的运动行为符合物理规律。渲染引擎提供逼真的视觉输出，'
  '支持光照、阴影和材质效果。传感器模型模拟摄像头、激光雷达等设备的成像过程，'
  '包括分辨率、视场角、噪声模型等参数均可配置。'
  'ros_gz_bridge将Gazebo内部的Protobuf消息自动转换为ROS 2标准消息格式，'
  '使得在仿真环境中开发的代码可直接迁移到实际机器人上运行。')

B('本课题使用TurtleBot3 waffle_pi作为仿真实验平台。'
  '该机器人配备RGB摄像头（640×480, 30Hz）和360度激光雷达，'
  '差速驱动底盘最大速度0.26 m/s，是验证室内避障算法的标准平台。'
  'TurtleBot3 World场景包含多面墙壁和立柱，构成了结构化的避障测试环境。')

H('2.3.3 nuScenes与Gazebo在本课题中的应用价值', 3)

B('nuScenes和Gazebo在本课题中承担互补的角色：'
  'nuScenes提供了大规模、高质量的真实世界驾驶图像，'
  '用于离线评估YOLO模型对各类障碍物的检测精度、推理速度和鲁棒性；'
  'Gazebo则提供了可控的虚拟仿真环境，'
  '用于在线测试完整的检测→决策→执行避障闭环。'
  '前者解决了"模型能检测什么"的问题，后者解决了"检测结果能否驱动安全避障"的问题。')

B('这种双数据源策略的优势在于：真实数据保证了评估结论的实际意义，'
  '仿真数据允许在受控条件下进行极端场景测试（如高速接近、多障碍物包围等），'
  '而这些场景在真实环境中难以安全复现。')

H('2.3.4 其他参考数据集简述', 3)

F(os.path.join(FIG, 'fig_2_7_dataset_comparison.png'),
  '图2-7 常用目标检测数据集与仿真平台对比')

B('图2-7对比了COCO、KITTI、nuScenes和Gazebo四种数据源的特性。'
  'COCO（Common Objects in Context）数据集包含33万张图像、80个日常目标类别，'
  '是YOLO系列模型预训练的标准数据源，本课题直接使用COCO预训练权重进行推理。'
  'KITTI数据集专注于自动驾驶场景，提供1.5万帧双目图像和LiDAR点云，'
  '虽然类别数较少（8类），但提供了精确的3D标注，常用于基准性能对比。'
  'nuScenes相比KITTI，场景规模更大、传感器配置更全面、类别标注更丰富，'
  '更适合评估多类别障碍物检测算法。Gazebo仿真可无限生成训练数据，'
  '且支持自动标注，在数据集扩展和极端场景测试方面具有独特优势。')

# ── 2.4 ──
H('2.4 模型轻量化与嵌入式部署技术', 2)

H('2.4.1 模型剪枝、量化、TensorRT加速原理', 3)

B('将深度学习模型部署到嵌入式设备上面临算力和内存的双重约束，'
  '需要通过模型轻量化技术在精度和效率之间取得平衡。'
  '常用的轻量化技术包括模型剪枝、权重量化和推理引擎优化。')

B('模型剪枝通过移除对输出贡献较小的网络参数（如接近零的权重或低重要性的通道）'
  '来减小模型规模。结构化剪枝直接删除整个卷积通道，'
  '可带来实际的推理加速而不依赖特殊硬件支持。')

B('权重量化将模型参数从FP32（32位浮点）降低到FP16（半精度）或INT8（8位整数），'
  '在大幅减少内存占用和计算量的同时，通过校准技术将精度损失控制在1%以内。'
  'FP16量化在现代GPU上可获得约2倍的计算加速，'
  'INT8量化可进一步获得约3～4倍的加速。')

F(os.path.join(FIG, 'fig_2_8_tensorrt_flow.png'),
  '图2-8 TensorRT模型优化与部署流程')

B('如图2-8所示，TensorRT是NVIDIA推出的高性能推理优化引擎，'
  '支持从PyTorch/ONNX模型到GPU推理引擎的端到端转换。'
  '其核心优化技术包括：（1）层融合——将多个连续的卷积、BN和激活层融合为单一操作，'
  '减少内存访问次数；（2）内核自动调优——为每个算子选择当前硬件上最快的CUDA内核实现；'
  '（3）精度校准——在FP16/INT8模式下通过校准数据集最小化量化误差。'
  '在本课题中，YOLO11s经TensorRT FP16优化后推理帧率从43.3 FPS提升至108.2 FPS，'
  '加速比达到2.5倍。')

H('2.4.2 Jetson Nano硬件特性与算力限制', 3)

F(os.path.join(FIG, 'fig_2_9_jetson_comparison.png'),
  '图2-9 嵌入式AI平台与开发平台硬件规格对比')

B('图2-9对比了Jetson Nano、Jetson Xavier NX和RTX 3090三个平台的硬件规格。'
  'Jetson Nano配备128核Maxwell GPU和4GB共享内存，AI算力为472 GFLOPS，'
  '功耗仅5～10W，是移动机器人部署的高性价比选择。'
  '其主要算力限制在于GPU核心数和内存带宽——128核Maxwell的计算能力约为'
  'RTX 3090的1/75，这意味着在Nano上运行YOLO模型必须依赖TensorRT的深度优化。'
  '经INT8量化后，YOLO11s可在Jetson Nano上达到约15 FPS，满足低速避障需求。')

B('Jetson Xavier NX作为中端方案，提供384核Volta GPU和21 TOPS的INT8算力，'
  '可在FP16模式下以约45 FPS运行YOLO11s，适合中高速避障场景。'
  '其价格约为Nano的4倍，需根据实际性能需求权衡选择。')

# ── 2.5 ──
H('2.5 本章小结', 2)

B('本章系统介绍了本课题涉及的四项关键技术：'
  '（1）YOLO目标检测算法从YOLOv1到YOLO11的演进历程和核心原理，'
  '以及YOLOv5与YOLO11在网络结构、检测头设计和训练策略方面的主要差异。'
  'YOLO的高速度、高精度和多规模档位设计使其成为移动机器人避障检测的理想选择。')

B('（2）ROS 2机器人操作系统基于DDS中间件的节点-话题-服务通信架构，'
  '为视觉感知与运动控制的模块化集成提供了标准化框架。'
  '话题通信机制天然适合检测→决策→执行的流水线数据处理。')

B('（3）nuScenes真实驾驶数据集和Gazebo仿真平台各具优势，'
  '前者提供多样化的真实场景用于离线评估，后者提供可控的虚拟环境用于在线闭环测试，'
  '二者互补使用可全面验证避障系统的性能。')

B('（4）TensorRT推理优化通过层融合、精度量化和内核调优等技术，'
  '可将YOLO模型的推理速度提升2.5～3.5倍，使其满足Jetson Nano等嵌入式平台的部署要求。'
  '这些技术基础为后续章节的模型设计、系统部署和实验验证奠定了理论支撑。')

doc.save(OUT_DOCX)
print(f"\n文档已保存: {OUT_DOCX}")
